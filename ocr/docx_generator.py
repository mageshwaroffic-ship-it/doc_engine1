"""
DOCX Generator: Creates an editable Word document with
the original scanned image as background and OCR text
overlaid at exact positions using DrawingML text boxes.
"""
from docx import Document
from docx.shared import Emu, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
# Register wps namespace for WordprocessingShape Since python-docx doesn't include it by default
nsmap['wps'] = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'

from PIL import Image, ImageOps
import pytesseract
import cv2
import numpy as np
import os


def preprocess_image(input_path: str, output_path: str) -> tuple:
    """
    Same preprocessing pipeline as main.py.
    Returns (preprocessed_path, margin_x_px, margin_y_px, orig_w, orig_h)
    so we can map coordinates back to the original image.
    """
    pimg = Image.open(input_path)
    pimg = ImageOps.exif_transpose(pimg)
    img = cv2.cvtColor(np.array(pimg), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape

    # Crop 8% borders
    margin_y = int(h * 0.08)
    margin_x = int(w * 0.08)
    gray = gray[margin_y:h - margin_y, margin_x:w - margin_x]

    # Shadow removal
    blurred_bg = cv2.GaussianBlur(gray, (0, 0), sigmaX=60)
    blurred_bg = np.where(blurred_bg == 0, 1, blurred_bg).astype(np.float32)
    normalized = (gray.astype(np.float32) / blurred_bg * 220)
    gray = np.clip(normalized, 0, 255).astype(np.uint8)

    # CLAHE
    clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
    gray = clahe.apply(gray)

    # Despeckle
    gray = cv2.medianBlur(gray, 3)

    # Adaptive threshold
    binary = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 41, 15)

    cv2.imwrite(output_path, binary)
    return output_path, margin_x, margin_y, w, h


def create_text_removed_image(original_image_path: str, preprocessed_path: str, output_path: str, margin_x: int, margin_y: int) -> str:
    """
    Create version of original image with text completely removed/erased.
    Uses inpainting to intelligently fill text regions with background color.
    """
    # Load original image
    pimg = Image.open(original_image_path)
    pimg = ImageOps.exif_transpose(pimg)
    orig_img = cv2.cvtColor(np.array(pimg), cv2.COLOR_RGB2BGR)
    
    # Load binary preprocessed image
    binary = cv2.imread(preprocessed_path, cv2.IMREAD_GRAYSCALE)
    
    # Crop binary to match the cropped region
    h, w = orig_img.shape[:2]
    binary_cropped = np.full((h, w), 255, dtype=np.uint8)
    if binary is not None:
        binary_cropped[margin_y:h-margin_y, margin_x:w-margin_x] = binary
    
    # Create mask where text is (white areas in binary = text = 255 in mask)
    text_mask = cv2.bitwise_not(binary_cropped)
    
    # Apply morphological operations to clean up the text mask
    kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    text_mask = cv2.dilate(text_mask, kernel, iterations=2)
    
    # Use inpainting to remove text (Telea algorithm)
    result = cv2.inpaint(orig_img, text_mask, 5, cv2.INPAINT_TELEA)
    
    # Save result
    cv2.imwrite(output_path, result)
    return output_path


def get_ocr_lines(preprocessed_path: str) -> list:
    """
    Run Tesseract with image_to_data() to get word-level bounding boxes.
    Group words into lines and return line data.
    """
    img = cv2.imread(preprocessed_path)
    data = pytesseract.image_to_data(
        img,
        config=r'-l kan+eng --oem 1 --psm 3 -c preserve_interword_spaces=1',
        output_type=pytesseract.Output.DICT
    )

    lines = {}
    n = len(data['text'])
    for i in range(n):
        text = data['text'][i].strip()
        if not text:
            continue
        conf = int(data['conf'][i])
        if conf < 0:
            continue
            
        kan_chars = sum(1 for c in text if 0x0C80 <= ord(c) <= 0x0CFF)
        eng_chars = sum(1 for c in text if c.lower() in 'abcdefghijklmnopqrstuvwxyz0123456789')
        if kan_chars == 0 and eng_chars == 0:
            continue

        line_key = (data['block_num'][i], data['par_num'][i], data['line_num'][i])
        if line_key not in lines:
            lines[line_key] = {
                'words': [],
                'left': data['left'][i],
                'right': data['left'][i] + data['width'][i],
                'tops': [],
                'heights': []
            }
        
        lines[line_key]['words'].append(text)
        
        # Expand horizontal bounding box
        lines[line_key]['left'] = min(lines[line_key]['left'], data['left'][i])
        lines[line_key]['right'] = max(lines[line_key]['right'], data['left'][i] + data['width'][i])
        
        # Keep track of tops and heights for median calculation
        lines[line_key]['tops'].append(data['top'][i])
        lines[line_key]['heights'].append(data['height'][i])
    
    import statistics
    
    # Convert to sorted list and use median for vertical positioning
    def get_robust_metric(values):
        if not values:
            return 0
        if len(values) < 3:
            return statistics.median(values)
        
        v_sorted = sorted(values)
        q1 = v_sorted[len(v_sorted)//4]
        q3 = v_sorted[3*len(v_sorted)//4]
        iqr = q3 - q1
        
        # Keep only values within 1.5 * IQR
        valid_values = [v for v in values if (q1 - 1.5 * iqr) <= v <= (q3 + 1.5 * iqr)]
        return statistics.median(valid_values if valid_values else values)

    result = []
    for key in sorted(lines.keys()):
        line = lines[key]
        
        robust_top = get_robust_metric(line['tops'])
        robust_height = get_robust_metric(line['heights'])
        
        result.append({
            'text': ' '.join(line['words']),
            'left': line['left'],
            'top': int(robust_top),
            'width': line['right'] - line['left'],
            'height': int(robust_height),
        })
    return result


def _make_background_image_anchor(rId, page_w_emu, page_h_emu):
    """Create a wp:anchor element for the background image (behind text)."""
    drawing = OxmlElement('w:drawing')
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '0')
    anchor.set('behindDoc', '1')       # BEHIND text
    anchor.set('locked', '1')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    simple_pos = OxmlElement('wp:simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')
    anchor.append(simple_pos)

    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'page')
    offset_h = OxmlElement('wp:posOffset')
    offset_h.text = '0'
    pos_h.append(offset_h)
    anchor.append(pos_h)

    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'page')
    offset_v = OxmlElement('wp:posOffset')
    offset_v.text = '0'
    pos_v.append(offset_v)
    anchor.append(pos_v)

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(page_w_emu))
    extent.set('cy', str(page_h_emu))
    anchor.append(extent)

    anchor.append(OxmlElement('wp:wrapNone'))

    doc_pr = OxmlElement('wp:docPr')
    doc_pr.set('id', '1')
    doc_pr.set('name', 'Background Image')
    anchor.append(doc_pr)

    # Graphic -> picture
    graphic = OxmlElement('a:graphic')
    graphic_data = OxmlElement('a:graphicData')
    graphic_data.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

    pic = OxmlElement('pic:pic')

    # nvPicPr
    nv = OxmlElement('pic:nvPicPr')
    cnv_pr = OxmlElement('pic:cNvPr')
    cnv_pr.set('id', '1')
    cnv_pr.set('name', 'background.jpg')
    nv.append(cnv_pr)
    nv.append(OxmlElement('pic:cNvPicPr'))
    pic.append(nv)

    # blipFill
    blip_fill = OxmlElement('pic:blipFill')
    blip = OxmlElement('a:blip')
    blip.set(qn('r:embed'), rId)
    blip_fill.append(blip)
    stretch = OxmlElement('a:stretch')
    stretch.append(OxmlElement('a:fillRect'))
    blip_fill.append(stretch)
    pic.append(blip_fill)

    # spPr
    sp_pr = OxmlElement('pic:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext = OxmlElement('a:ext')
    ext.set('cx', str(page_w_emu))
    ext.set('cy', str(page_h_emu))
    xfrm.append(ext)
    sp_pr.append(xfrm)
    prst = OxmlElement('a:prstGeom')
    prst.set('prst', 'rect')
    prst.append(OxmlElement('a:avLst'))
    sp_pr.append(prst)
    pic.append(sp_pr)

    graphic_data.append(pic)
    graphic.append(graphic_data)
    anchor.append(graphic)
    drawing.append(anchor)
    return drawing


def _make_text_box_anchor(box_id, text, x_emu, y_emu, w_emu, h_emu, font_size_half_pt):
    """Create a wp:anchor element for an editable transparent text box."""
    drawing = OxmlElement('w:drawing')
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', str(box_id + 10))
    anchor.set('behindDoc', '0')       # IN FRONT of background
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    simple_pos = OxmlElement('wp:simplePos')
    simple_pos.set('x', '0')
    simple_pos.set('y', '0')
    anchor.append(simple_pos)

    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'page')
    offset_h = OxmlElement('wp:posOffset')
    offset_h.text = str(x_emu)
    pos_h.append(offset_h)
    anchor.append(pos_h)

    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'page')
    offset_v = OxmlElement('wp:posOffset')
    offset_v.text = str(y_emu)
    pos_v.append(offset_v)
    anchor.append(pos_v)

    extent = OxmlElement('wp:extent')
    extent.set('cx', str(w_emu))
    extent.set('cy', str(h_emu))
    anchor.append(extent)

    anchor.append(OxmlElement('wp:wrapNone'))

    doc_pr = OxmlElement('wp:docPr')
    doc_pr.set('id', str(box_id))
    doc_pr.set('name', f'TextBox {box_id}')
    anchor.append(doc_pr)

    # Graphic -> wordprocessingShape text box
    graphic = OxmlElement('a:graphic')
    graphic_data = OxmlElement('a:graphicData')
    graphic_data.set('uri', 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')

    wsp = OxmlElement('wps:wsp')

    cnv = OxmlElement('wps:cNvSpPr')
    cnv.set('txBox', '1')
    wsp.append(cnv)

    # Shape properties (transparent, no border)
    sp_pr = OxmlElement('wps:spPr')
    xfrm = OxmlElement('a:xfrm')
    off = OxmlElement('a:off')
    off.set('x', '0')
    off.set('y', '0')
    xfrm.append(off)
    ext = OxmlElement('a:ext')
    ext.set('cx', str(w_emu))
    ext.set('cy', str(h_emu))
    xfrm.append(ext)
    sp_pr.append(xfrm)

    prst = OxmlElement('a:prstGeom')
    prst.set('prst', 'rect')
    prst.append(OxmlElement('a:avLst'))
    sp_pr.append(prst)

    sp_pr.append(OxmlElement('a:noFill'))
    ln = OxmlElement('a:ln')
    ln.set('w', '0')
    ln.append(OxmlElement('a:noFill'))
    sp_pr.append(ln)
    wsp.append(sp_pr)

    # Text box content
    txbx = OxmlElement('wps:txbx')
    txbx_content = OxmlElement('w:txbxContent')
    p = OxmlElement('w:p')

    # Paragraph properties
    ppr = OxmlElement('w:pPr')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    ppr.append(spacing)
    p.append(ppr)

    # Run with text
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size_half_pt))
    rpr.append(sz)
    sz_cs = OxmlElement('w:szCs')
    sz_cs.set(qn('w:val'), str(font_size_half_pt))
    rpr.append(sz_cs)

    # Transparent text color (very light, almost invisible so bg image shows)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '333333')
    rpr.append(color)

    # Detect if Kannada text and set appropriate font
    has_kannada = any(0x0C80 <= ord(c) <= 0x0CFF for c in text)
    r_fonts = OxmlElement('w:rFonts')
    if has_kannada:
        r_fonts.set(qn('w:ascii'), 'Noto Sans Kannada')
        r_fonts.set(qn('w:hAnsi'), 'Noto Sans Kannada')
        r_fonts.set(qn('w:cs'), 'Noto Sans Kannada')
    else:
        r_fonts.set(qn('w:ascii'), 'Noto Sans')
        r_fonts.set(qn('w:hAnsi'), 'Noto Sans')
        r_fonts.set(qn('w:cs'), 'Noto Sans')
    rpr.append(r_fonts)

    r.append(rpr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)

    txbx_content.append(p)
    txbx.append(txbx_content)
    wsp.append(txbx)

    # Body properties - prevent wrapping
    body_pr = OxmlElement('wps:bodyPr')
    body_pr.set('wrap', 'none')
    body_pr.set('lIns', '0')
    body_pr.set('tIns', '0')
    body_pr.set('rIns', '0')
    body_pr.set('bIns', '0')
    body_pr.set('anchor', 't')
    wsp.append(body_pr)

    graphic_data.append(wsp)
    graphic.append(graphic_data)
    anchor.append(graphic)
    drawing.append(anchor)
    return drawing


def generate_formatted_docx(original_image_path: str, output_docx_path: str, job_dir: str, edited_text_data: list = None):
    """
    Main function: Creates a .docx with:
    1. Original image with text regions blacked out (masked)
    2. OCR text overlaid as editable text boxes at exact positions
    """
    # 1. Preprocess the image to get text positions
    preprocessed_path = os.path.join(job_dir, "preprocessed.png")
    _, margin_x, margin_y, orig_w, orig_h = preprocess_image(original_image_path, preprocessed_path)

    # 2. Create image with text completely removed/erased
    text_removed_path = os.path.join(job_dir, "text_removed_bg.jpg")
    create_text_removed_image(original_image_path, preprocessed_path, text_removed_path, margin_x, margin_y)
    print(f"  Text removed from background image (using inpainting)")

    # 3. Get OCR line positions from preprocessed image (or use edited text if provided)
    if edited_text_data:
        lines = edited_text_data
        print(f"  Using {len(lines)} edited text items from frontend")
    else:
        lines = get_ocr_lines(preprocessed_path)
        print(f"  Found {len(lines)} text lines from OCR")

    # 4. Create DOCX document
    doc = Document()
    section = doc.sections[0]

    # Set page size to match image aspect ratio (fit in ~11 inch height)
    aspect = orig_w / orig_h
    if aspect > 1:  # landscape
        page_w_emu = int(11 * 914400)
        page_h_emu = int(page_w_emu / aspect)
    else:  # portrait
        page_h_emu = int(11 * 914400)
        page_w_emu = int(page_h_emu * aspect)

    section.page_width = Emu(page_w_emu)
    section.page_height = Emu(page_h_emu)
    section.left_margin = Emu(0)
    section.right_margin = Emu(0)
    section.top_margin = Emu(0)
    section.bottom_margin = Emu(0)

    # 5. Calculate scale factors
    # Preprocessed image coordinates -> original image coordinates -> page EMU
    scale_x = page_w_emu / orig_w   # EMU per original pixel
    scale_y = page_h_emu / orig_h

    # Offset for the 8% crop
    offset_x_emu = int(margin_x * scale_x)
    offset_y_emu = int(margin_y * scale_y)

    # Scale for preprocessed pixels (preprocessed image is 84% of original)
    prep_w = orig_w - 2 * margin_x
    prep_h = orig_h - 2 * margin_y
    prep_scale_x = (page_w_emu - 2 * offset_x_emu) / prep_w if prep_w > 0 else scale_x
    prep_scale_y = (page_h_emu - 2 * offset_y_emu) / prep_h if prep_h > 0 else scale_y

    # 6. Safely load the TEXT-REMOVED image into the package and get its relationship ID
    para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    inline_shape = para.add_run().add_picture(text_removed_path)
    blip = inline_shape._inline.xpath('.//a:blip')[0]
    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
    
    # Remove the dummy inline picture run now that the image is safely loaded
    para.clear()

    # Clear default paragraph spacing
    ppr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    ppr.append(spacing)

    # Add text-removed background image run
    bg_run = OxmlElement('w:r')
    bg_drawing = _make_background_image_anchor(rId, page_w_emu, page_h_emu)
    bg_run.append(bg_drawing)
    para._p.append(bg_run)

    # 7. Pre-process lines: Sort and Resolve Overlaps
    # Ensure lines are processed in vertical order
    if 'top' in lines[0] or 'y' in lines[0]:
        lines.sort(key=lambda l: l.get('top', l.get('y', 0)))
    
    # Track the last used bottom position (in points or EMU) to prevent collision
    last_bottom_emu = 0
    
    box_id = 2
    for line in lines:
        text = line.get('text', line.get('text', ''))
        if not text or not text.strip():
            continue

        # Handle both OCR lines and edited text data
        if 'left' in line:
            # OCR line format: left, top, width, height (preprocessed coords)
            x_emu = int(line['left'] * prep_scale_x + offset_x_emu)
            y_emu = int(line['top'] * prep_scale_y + offset_y_emu)
            w_emu = int(line['width'] * prep_scale_x)
            h_emu = int(line['height'] * prep_scale_y)
        else:
            # Edited text format: x, y, width, height (original image coords)
            x_emu = int(line['x'] * scale_x)
            y_emu = int(line['y'] * scale_y)
            w_emu = int(line['width'] * scale_x)
            h_emu = int(line['height'] * scale_y)

        # COLLISION PREVENTION: If this box overlaps the previous one, push it down
        # Increased gap to 4 pixels (in EMU) for better vertical breathing room
        if y_emu < last_bottom_emu + Emu(4).emu:
            y_emu = last_bottom_emu + Emu(4).emu

        # Use ACTUAL height from OCR as the font size (don't calculate/predict)
        # Box height in pixels → scale to EMU → convert to points
        h_emu_for_font = h_emu if 'left' in line else int(line['height'] * scale_y)
        
        # Convert EMU height to points for font size
        # Apply 0.8x safety factor to prevent font line-height from causing overlaps
        # (A font size of X pt usually results in a line height > X pt)
        font_size_pt = max(8, (h_emu_for_font / 12700) * 0.8) 
        # 7.1. Auto-shrink font if text is too wide for the box
        # Character width heuristic (average Sans Serif / Kannada)
        char_factor = 0.65 if has_kannada else 0.55
        est_width_pt = len(text) * font_size_pt * char_factor
        box_width_pt = w_emu / 12700
        
        if est_width_pt > box_width_pt:
            shrink_factor = (box_width_pt / est_width_pt) * 0.95 # Extra 5% breathing room
            font_size_pt = font_size_pt * shrink_factor
            font_size_half_pt = int(font_size_pt * 2)

        # Final sanity bound for half-points (clamp to max 48pt for titles)
        font_size_half_pt = max(12, min(font_size_half_pt, 96)) 
        
        # Update last_bottom to track collision
        last_bottom_emu = y_emu + h_emu

        run = OxmlElement('w:r')
        text_box = _make_text_box_anchor(box_id, text, x_emu, y_emu, w_emu, h_emu, font_size_half_pt)
        run.append(text_box)
        para._p.append(run)
        box_id += 1

    # 8. Save DOCX
    doc.save(output_docx_path)
    print(f"  Saved DOCX: {output_docx_path}")
    return output_docx_path
